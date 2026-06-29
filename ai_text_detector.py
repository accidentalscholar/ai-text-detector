# ==========================================
# AI Text Detector 
# Version: 23.0
# Citation: Pundir, V. (2026, June 29). AI Text Detector Version (23.0). Retrieved from https://github.com/accidentalscholar/ai-text-detector. 
# Citation: RIS and BibTeX files included for referencing software.
# Tested in: Python 3.10.9 64 bit packaged by Anaconda, Inc.
# Reporsitory: https://github.com/accidentalscholar/ai-text-detector
# Provided under: GNU AFFERO GENERAL PUBLIC LICENSE (see accompanying license file)
# ==========================================

import subprocess
import sys
import os
import gc
import importlib.util
from datetime import datetime

# Suppress Hugging Face Symlink Warnings on Windows
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# --- 1. CUSTOM PATH CONFIGURATION ---
def load_custom_paths():
    """Reads path.txt and appends directories to system and python paths."""
    path_file = "path.txt"
    if os.path.exists(path_file):
        if os.path.getsize(path_file) == 0:
            print(f"Found {path_file} but it is empty. Moving on to default system paths.")
            return
            
        print(f"Found {path_file}. Loading custom paths...")
        with open(path_file, "r") as f:
            for line in f:
                custom_path = line.strip()
                # Check if it's a valid directory and not already empty
                if custom_path and os.path.isdir(custom_path):
                    print(f" -> Adding to paths: {custom_path}")
                    # 1. Add to Python path (for libraries) - APPEND is safer to avoid shadowing system libs
                    if custom_path not in sys.path:
                        sys.path.append(custom_path)
                    # 2. Add to System path (for compiled executables)
                    if custom_path not in os.environ["PATH"]:
                        os.environ["PATH"] = custom_path + os.pathsep + os.environ["PATH"]
    else:
        print(f"No {path_file} found. Relying on default system paths.")

# Execute path loading before importing external libraries
load_custom_paths()


# --- 2. FAULT TOLERANCE: SILENT AUTO-INSTALL ---
def ensure_library(package_name, import_name=None):
    """
    Attempts to find a library using importlib. If missing, attempts silent installation.
    Using find_spec prevents circular import bugs in Python 3.13 for massive libraries like PyTorch.
    """
    if import_name is None:
        import_name = package_name
    
    # Safely check if the library is installed WITHOUT executing its __init__.py code
    if importlib.util.find_spec(import_name) is not None:
        return True
        
    print(f"[*] Dependency '{package_name}' not found. Attempting to install it now...")
    try:
        # DEVNULL suppresses the standard pip output and error text
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", package_name],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        print(f"Successfully installed {package_name}.")
        return True
    except subprocess.CalledProcessError:
        print("\n" + "="*50)
        print(" INSTALLATION FAILED ")
        print("="*50)
        print(f"Could not automatically install '{package_name}'.")
        print(f"Please open your Spyder console or terminal and run:")
        print(f"\n    pip install {package_name}\n")
        print("="*50 + "\n")
        return False

# Ensure all required libraries are present; stop execution gracefully if one fails
dependencies = [
    ("transformers", None),
    ("torch", None),
    ("python-docx", "docx"),
    ("pypdf", None),
    ("pandas", None),
    ("XlsxWriter", "xlsxwriter"),
    ("sentencepiece", None),
    ("protobuf", None),
    ("psutil", None) 
]

for pkg, imp in dependencies:
    if not ensure_library(pkg, imp):
        print("Dependency check failed. Exiting script gracefully.")
        # Raise a custom exception rather than sys.exit() to prevent Spyder kernel crashes
        raise SystemExit("Missing Dependencies")


# --- 3. CORE IMPORTS ---
import tkinter as tk
from tkinter import filedialog
import docx
from pypdf import PdfReader
import pandas as pd
import psutil

# --- SPYDER-SPECIFIC FIX FOR PYTORCH ---
# If an import fails in Spyder, it leaves a "partially initialized" broken module 
# in memory. This block detects that corruption, clears it, and retries.
try:
    import torch
except AttributeError as e:
    if "partially initialized" in str(e) or "circular import" in str(e):
        print("\n[*] Detected a stuck/corrupted PyTorch module in Spyder's memory. Clearing it...")
        for key in list(sys.modules.keys()):
            if key == 'torch' or key.startswith('torch.'):
                del sys.modules[key]
        import torch # Retry fresh import
    else:
        raise

from transformers import pipeline
import numpy as np


# --- 4. HARDWARE ACCELERATION DETECTION ---
def get_optimal_device():
    """Detects if an NVIDIA GPU or Apple Silicon is available."""
    if torch.cuda.is_available():
        print("Hardware Acceleration: NVIDIA GPU (CUDA) detected.")
        return torch.device("cuda"), 0 
    elif hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
        print("Hardware Acceleration: Apple Silicon (MPS) detected.")
        return torch.device("mps"), "mps"
    else:
        print("Hardware Acceleration: None. Defaulting to CPU.")
        return torch.device("cpu"), -1

TORCH_DEVICE, PIPELINE_DEVICE = get_optimal_device()


# --- 5. HELPER FUNCTIONS ---
def get_text_from_docx(file_path):
    """Safely extracts text from a docx file."""
    try:
        doc = docx.Document(file_path)
        return " ".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return ""

def get_text_from_pdf(file_path):
    """Safely extracts text from a pdf file."""
    try:
        reader = PdfReader(file_path)
        text_chunks = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text_chunks.append(extracted.strip())
        return " ".join(text_chunks)
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return ""

def chunk_generator(text, max_words=300):
    """Yields text chunks one at a time to keep memory flat."""
    words = text.split()
    for i in range(0, len(words), max_words):
        yield ' '.join(words[i:i + max_words])

def select_folder():
    """Opens a GUI dialog to select a folder."""
    root = tk.Tk()
    root.withdraw() 
    
    # Force the dialog to appear in the foreground
    root.attributes('-topmost', True)
    root.lift()
    
    folder_path = filedialog.askdirectory(parent=root, title="Select Folder Containing DOCX/PDF Files")
    root.destroy() # Clean up the root window after selection
    return folder_path

def clear_memory():
    """Aggressively frees up RAM/VRAM after a model is done."""
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()


# --- 6. MAIN SEQUENTIAL PIPELINE ---
def main():
    folder_path = select_folder()
    if not folder_path:
        print("No folder selected. Exiting.")
        return

    # Phase 1: Pre-read all documents
    documents = {} 
    print(f"Scanning and extracting text from: {folder_path}")
    
    for filename in os.listdir(folder_path):
        if filename.startswith("~"):
            continue
            
        file_path = os.path.join(folder_path, filename)
        text = ""
        
        # Route to appropriate extractor based on extension
        if filename.lower().endswith(".docx"):
            text = get_text_from_docx(file_path)
        elif filename.lower().endswith(".pdf"):
            text = get_text_from_pdf(file_path)
        else:
            continue # Skip unsupported files

        if text:
            documents[filename] = text
        else:
            print(f"Skipping {filename} (Empty or unreadable)")

    if not documents:
        print("No valid DOCX or PDF data extracted.")
        return

    results = {filename: {} for filename in documents.keys()}

    # ---------------------------------------------------------
    # PHASE 2: DYNAMIC PRE-TRAINED MODELS ENSEMBLE
    # ---------------------------------------------------------
    # Grouped by Architecture for clean Excel columns
    DETECTION_MODELS = [
        # --- RoBERTa Architecture Group ---
        {
            "repo": "roberta-base-openai-detector",
            "name": "RoBERTa: OpenAI Base",
            "ai_labels": ["fake"]
        },
        {
            "repo": "roberta-large-openai-detector",
            "name": "RoBERTa: OpenAI Large",
            "ai_labels": ["fake"]
        },
        {
            "repo": "Hello-SimpleAI/chatgpt-detector-roberta",
            "name": "RoBERTa: ChatGPT Detector",
            "ai_labels": ["chatgpt"]
        },
        # --- DeBERTa Architecture Group ---
        {
            "repo": "ogmatrixllm/glyph-v1.1",
            "name": "DeBERTa: GLYPH v1.1",
            "ai_labels": ["label_1", "1", "ai", "generated", "fake"]
        },
        # --- ModernBERT Architecture Group ---
        {
            "repo": "AICodexLab/answerdotai-ModernBERT-base-ai-detector",
            "name": "ModernBERT: AI Codex",
            "ai_labels": ["label_1", "1", "ai", "generated", "fake"]
        },
        # --- Diverse/Mixed Architecture Group ---
        {
            "repo": "PirateXX/AI-Content-Detector",
            "name": "Diverse: PirateXX Detector",
            "ai_labels": ["label_0"] # Correctly mapped to LABEL_0
        }
    ]

    for idx, model_config in enumerate(DETECTION_MODELS, 1):
        print(f"\n--- Starting Phase {idx}: {model_config['name']} ---")
        print(f"Loading {model_config['repo']} from Hugging Face...")
        
        try:
            hf_classifier = pipeline(
                "text-classification", 
                model=model_config['repo'], 
                truncation=True, 
                max_length=512,
                device=PIPELINE_DEVICE
            )
            
            for filename, text in documents.items():
                scores = []
                for chunk in chunk_generator(text):
                    res = hf_classifier(chunk)[0]
                    
                    # Robust label checker: True if the model's output label implies AI
                    raw_label = str(res['label']).lower()
                    
                    # Check exact matches against our trigger words (case-insensitive)
                    is_ai = any(trigger.lower() == raw_label for trigger in model_config['ai_labels'])
                    
                    # Calculate final AI probability
                    score = res['score'] if is_ai else 1.0 - res['score']
                    scores.append(score)
                
                results[filename][model_config['name']] = round(sum(scores) / len(scores), 4) if scores else 0
                print(f"[{filename}] {model_config['name']} Complete.")

            # Unload model to free up RAM/VRAM before the next one starts
            del hf_classifier
            clear_memory()
            
        except Exception as e:
            print(f"Error loading or running {model_config['repo']}: {e}")
            for filename in documents.keys():
                results[filename][model_config['name']] = 0

    
    # ---------------------------------------------------------
    # PHASE 3: OPTIONAL CHECKFOR.AI (PANGRAM) INTEGRATION
    # ---------------------------------------------------------
    # Find script directory reliably even if running inside Spyder's IPython console
    script_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
    editlens_token_path = os.path.join(script_dir, "editlens_token.txt")

    if os.path.exists(editlens_token_path):
        print("\n--- Starting Optional Phase: Pangram Labs (formerly Checkfor.ai) ---")
        try:
            with open(editlens_token_path, "r") as f:
                editlens_token = f.read().strip()
                
            if editlens_token:
                print("Found editlens_token.txt. Authenticating with Hugging Face...")
                pangram_repo = "pangram/editlens_roberta-large"
                print(f"Loading {pangram_repo} (Gated Model)...")
                
                try:
                    pangram_classifier = pipeline(
                        "text-classification", 
                        model=pangram_repo, 
                        token=editlens_token,
                        truncation=True, 
                        max_length=512,
                        device=PIPELINE_DEVICE
                    )
                    
                    for filename, text in documents.items():
                        scores = []
                        for chunk in chunk_generator(text):
                            res = pangram_classifier(chunk)[0]
                            raw_label = str(res['label']).lower()
                            
                            # Pangram uses ternary/binary EditLens classification. We trigger on any AI keywords.
                            ai_triggers = ["ai_generated", "ai_edited", "ai", "label_1", "1", "fake", "machine"]
                            is_ai = any(trigger in raw_label for trigger in ai_triggers)
                            
                            score = res['score'] if is_ai else 1.0 - res['score']
                            scores.append(score)
                        
                        results[filename]["RoBERTa: Pangram EditLens"] = round(sum(scores) / len(scores), 4) if scores else 0
                        print(f"[{filename}] Pangram EditLens Complete.")
                    
                    del pangram_classifier
                    clear_memory()
                    print("Pangram EditLens execution successful.")
                    
                except Exception as e:
                    print(f"Failed to run Pangram EditLens. Ensure your token is valid and you have accepted the model terms on Hugging Face. Error: {e}")
                    # Remove incomplete data so the Excel sheet stays clean
                    for filename in documents.keys():
                        results[filename].pop("RoBERTa: Pangram EditLens", None)
            else:
                print("editlens_token.txt is empty. Skipping Pangram (Checkfor.ai) phase.")
        except Exception as e:
            print(f"Error reading editlens_token.txt: {e}")


    # ---------------------------------------------------------
    # PHASE 4: OPTIONAL DMICZ BINOCULARS INTEGRATION
    # ---------------------------------------------------------
    bino_file_path = os.path.join(script_dir, "binoculars.txt")

    if os.path.exists(bino_file_path):
        print("\n--- Starting Optional Phase: Dmicz's Binoculars (Zero-Shot) ---")
        try:
            if os.path.getsize(bino_file_path) == 0 or open(bino_file_path).read().strip() == "":
                print("Found binoculars.txt trigger file. Ascertaining system resources...")
                
                # Retrieve System RAM and GPU VRAM 
                sys_ram_gb = psutil.virtual_memory().total / (1024**3)
                gpu_vram_gb = 0.0
                if torch.cuda.is_available():
                    gpu_vram_gb = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                
                # Retrieve Free Disk Space from the Hugging Face cache location
                hf_cache_dir = os.environ.get("HF_HOME", os.path.expanduser("~"))
                try:
                    disk_usage = psutil.disk_usage(hf_cache_dir)
                    free_space_gb = disk_usage.free / (1024**3)
                except Exception:
                    # Fallback to root/C: drive if path mapping fails
                    disk_usage = psutil.disk_usage(os.path.abspath(os.sep))
                    free_space_gb = disk_usage.free / (1024**3)
                
                # Binoculars loads 2x Falcon-7B models (~30GB total). 
                REQUIRED_DISK_GB = 60.0
                
                # Needs >= 16GB VRAM or >= 32GB System RAM, AND >= 60GB Free Disk
                if (gpu_vram_gb < 16.0 and sys_ram_gb < 32.0) or free_space_gb < REQUIRED_DISK_GB:
                    msg = (f"Binoculars skipped: Insufficient resources.\n"
                           f"Found: {gpu_vram_gb:.1f}GB VRAM, {sys_ram_gb:.1f}GB System RAM, {free_space_gb:.1f}GB Free Disk Space.\n"
                           f"Requires: (>= 16.0GB VRAM OR >= 32.0GB System RAM) AND >= {REQUIRED_DISK_GB}GB Free Disk Space.")
                    print(msg)
                    with open(bino_file_path, "w") as f:
                        f.write(msg)
                else:
                    print(f"Resources sufficient ({free_space_gb:.1f}GB disk space found). Preparing Binoculars...")
                    
                    if not ensure_library("git+https://github.com/ahans30/Binoculars.git", "binoculars"):
                        msg = "Binoculars skipped: Failed to install binoculars package via pip. Ensure 'git' developer tools are installed on Windows."
                        print(msg)
                        with open(bino_file_path, "w") as f:
                            f.write(msg)
                    else:
                        print("Downloading models and initializing Zero-Shot Analysis... (This may take a long time)")
                        try:
                            from binoculars import Binoculars
                            bino = Binoculars()
                            
                            for filename in documents.keys():
                                results[filename]["Diverse: Binoculars"] = 0
                                
                            for filename, text in documents.items():
                                scores = []
                                for chunk in chunk_generator(text):
                                    pred = bino.predict(chunk)
                                    if isinstance(pred, (list, tuple)):
                                        pred = pred[0]
                                    
                                    pred_str = str(pred).lower()
                                    if "ai" in pred_str or "machine" in pred_str or "fake" in pred_str:
                                        scores.append(1.0)
                                    else:
                                        scores.append(0.0)
                                        
                                results[filename]["Diverse: Binoculars"] = round(sum(scores) / len(scores), 4) if scores else 0
                                print(f"[{filename}] Binoculars Complete.")
                            
                            del bino
                            clear_memory()
                            print("Binoculars execution successful.")
                            with open(bino_file_path, "w") as f:
                                f.write("SUCCESS: Binoculars ran and completed on this execution.")
                                
                        except Exception as e:
                            for filename in documents.keys():
                                results[filename].pop("Diverse: Binoculars", None)
                            msg = f"Binoculars skipped due to fatal runtime error: {e}"
                            print(msg)
                            with open(bino_file_path, "w") as f:
                                f.write(msg)
            else:
                print("binoculars.txt is NOT empty. Skipping Binoculars phase.")
        except Exception as e:
            print(f"Error reading binoculars.txt: {e}")


    # ---------------------------------------------------------
    # PHASE 5: EXPORT TO EXCEL
    # ---------------------------------------------------------
    print("\nCompiling Data into Excel...")
    
    # Generate timestamp for filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"LLM_Detection_Results_{timestamp}.xlsx"
    output_file = os.path.join(folder_path, output_filename)
    
    final_data = []
    for filename, data in results.items():
        row = {"Filename": filename}
        row.update(data)
        final_data.append(row)

    df = pd.DataFrame(final_data)
    writer = pd.ExcelWriter(output_file, engine='xlsxwriter')
    
    # Write data starting on row 1 (skipping pandas headers) to make room for Excel Table headers
    df.to_excel(writer, sheet_name='Detection Results', index=False, header=False, startrow=1)
    
    workbook  = writer.book
    worksheet = writer.sheets['Detection Results']
    
    num_rows = len(df)
    num_cols = len(df.columns)
    
    # Add native Excel Table formatting dynamically based on the number of models ran
    column_settings = [{'header': column} for column in df.columns]
    worksheet.add_table(0, 0, num_rows, num_cols - 1, {'columns': column_settings})

    # Define formats
    percent_format = workbook.add_format({'num_format': '0.00%'})
    format_yellow = workbook.add_format({'bg_color': '#FFFF99'}) 
    format_orange = workbook.add_format({'bg_color': '#FFCC99'}) 
    format_red    = workbook.add_format({'bg_color': '#FF9999'}) 

    # Apply conditional formatting dynamically to all model score columns
    for col_num in range(1, num_cols):
        col_letter = chr(65 + col_num)
        cell_range = f'{col_letter}2:{col_letter}{num_rows + 1}'

        worksheet.conditional_format(cell_range, {'type': 'cell', 'criteria': '>', 'value': 0.5, 'format': format_red})
        worksheet.conditional_format(cell_range, {'type': 'cell', 'criteria': 'between', 'minimum': 0.3, 'maximum': 0.5, 'format': format_orange})
        worksheet.conditional_format(cell_range, {'type': 'cell', 'criteria': 'between', 'minimum': 0.1, 'maximum': 0.2999, 'format': format_yellow})

    worksheet.set_column('A:A', 40)
    
    # Apply width and the percent format simultaneously to all dynamic score columns
    # Calculate the last column letter using chr() arithmetic (works up to column Z)
    last_col_letter = chr(65 + num_cols - 1)
    worksheet.set_column(f'B:{last_col_letter}', 25, percent_format)

    writer.close()
    print(f"\nProcess complete! Saved to: {output_file}")

if __name__ == "__main__":
    try:
        main()
    except SystemExit as e:
        print(e)